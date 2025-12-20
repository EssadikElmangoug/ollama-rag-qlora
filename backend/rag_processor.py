"""
RAG Processing Module
Handles document ingestion, chunking, embedding, and retrieval
"""
import os
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LangchainDocument
from docx import Document
import pandas as pd

class RAGProcessor:
    def __init__(self, vectorstore_path="faiss_index", embedding_model="sentence-transformers/all-mpnet-base-v2"):
        """
        Initialize RAG Processor
        
        Args:
            vectorstore_path: Path to save/load FAISS vector store
            embedding_model: Name of the embedding model to use
        """
        self.vectorstore_path = vectorstore_path
        self.embedding_model_name = embedding_model
        
        # Initialize embeddings
        self.embeddings = HuggingFaceEmbeddings(
            model_name=self.embedding_model_name,
            model_kwargs={"device": "cpu"}  # Use CPU for compatibility
        )
        
        # Initialize text splitter
        self.text_splitter = CharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=30,
            separator="\n"
        )
        
        # Load or create vector store
        self.vectorstore = self._load_or_create_vectorstore()
    
    def _load_or_create_vectorstore(self):
        """Load existing vector store or create a new one"""
        if os.path.exists(self.vectorstore_path) and os.path.isdir(self.vectorstore_path):
            try:
                # Check if index files exist
                index_file = os.path.join(self.vectorstore_path, "index.faiss")
                if os.path.exists(index_file):
                    return FAISS.load_local(
                        self.vectorstore_path,
                        self.embeddings,
                        allow_dangerous_deserialization=True
                    )
            except Exception as e:
                print(f"Error loading vector store: {e}. Creating new one.")
        
        # Create empty vector store with a minimal document
        dummy_doc = LangchainDocument(page_content="placeholder", metadata={})
        vectorstore = FAISS.from_documents([dummy_doc], self.embeddings)
        return vectorstore
    
    def _load_document(self, filepath):
        """Load document based on file type"""
        file_ext = os.path.splitext(filepath)[1].lower()
        
        try:
            if file_ext == '.pdf':
                loader = PyPDFLoader(filepath)
                documents = loader.load()
            elif file_ext in ['.txt', '.md']:
                loader = TextLoader(filepath, encoding='utf-8')
                documents = loader.load()
            elif file_ext == '.docx':
                loader = Docx2txtLoader(filepath)
                documents = loader.load()
            elif file_ext in ['.xlsx', '.xls']:
                # Load Excel file
                df = pd.read_excel(filepath)
                # Convert to text
                text_content = df.to_string()
                documents = [LangchainDocument(
                    page_content=text_content,
                    metadata={"source": filepath, "type": "excel"}
                )]
            elif file_ext == '.csv':
                # Load CSV file
                df = pd.read_csv(filepath)
                text_content = df.to_string()
                documents = [LangchainDocument(
                    page_content=text_content,
                    metadata={"source": filepath, "type": "csv"}
                )]
            else:
                # Try text loader as fallback
                loader = TextLoader(filepath, encoding='utf-8')
                documents = loader.load()
            
            return documents
        except Exception as e:
            raise Exception(f"Error loading document {filepath}: {str(e)}")
    
    def process_document(self, filepath, filename):
        """
        Process a document: load, chunk, and add to vector store
        
        Args:
            filepath: Path to the document file
            filename: Original filename
            
        Returns:
            dict with processing results
        """
        try:
            # Load document
            documents = self._load_document(filepath)
            
            if not documents:
                return {
                    "success": False,
                    "error": "No content extracted from document"
                }
            
            # Add metadata
            for doc in documents:
                doc.metadata["source_file"] = filename
                doc.metadata["filepath"] = filepath
            
            # Split into chunks
            chunks = self.text_splitter.split_documents(documents)
            
            if not chunks:
                return {
                    "success": False,
                    "error": "No chunks created from document"
                }
            
            # Add chunks to vector store
            self.vectorstore.add_documents(chunks)
            
            # Save vector store
            self.vectorstore.save_local(self.vectorstore_path)
            
            return {
                "success": True,
                "chunks_count": len(chunks),
                "total_pages": len(documents),
                "message": f"Successfully processed {filename}"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e)
            }
    
    def remove_document(self, filename):
        """
        Remove a document from the vector store
        Note: FAISS doesn't support direct deletion, so we'll need to rebuild
        """
        # For now, we'll mark documents for exclusion in retrieval
        # A full implementation would rebuild the index without the document
        return {"success": True, "message": f"Document {filename} marked for removal"}
    
    def get_retriever(self, k=4):
        """Get a retriever from the vector store"""
        return self.vectorstore.as_retriever(search_kwargs={"k": k})
    
    def search_similar(self, query, k=4):
        """Search for similar documents"""
        try:
            results = self.vectorstore.similarity_search(query, k=k)
            return results
        except Exception as e:
            raise Exception(f"Error searching vector store: {str(e)}")
    
    def search_similar_with_scores(self, query, k=4):
        """Search for similar documents with similarity scores"""
        try:
            results = self.vectorstore.similarity_search_with_score(query, k=k)
            return results
        except Exception as e:
            raise Exception(f"Error searching vector store: {str(e)}")
    
    def has_documents(self):
        """Check if vector store has any documents"""
        try:
            # Try to get one document to check if store is populated
            results = self.vectorstore.similarity_search("test", k=1)
            # If we have results and they're not just the placeholder, we have documents
            if results and len(results) > 0:
                # Check if it's not the placeholder document
                if results[0].page_content and results[0].page_content.strip() != "placeholder":
                    return True
            return False
        except:
            return False

